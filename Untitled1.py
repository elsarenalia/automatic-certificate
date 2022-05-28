#!/usr/bin/env python
# coding: utf-8

# In[1]:


import pandas as pd
from docx import Document


# In[2]:


#NamaTemplate=input("Masukan nama file template beserta eksistensinya")
NamaTemplate='00 temp sertifikat cirebon.docx'

#FileData=input("Masukan nama file data beserta eksistensinya")
NamaData='penerima e-sertifikat.xlsx'

#load data hingga menjadi tuple
kotor=pd.read_excel(NamaData, index_col=0)
mid = kotor.to_records(index=False) #to_records : Convert DataFrame to a NumPy record array.
data = tuple(mid)


# In[3]:


def mengisi(d,idx,NamaTemplate):
    document = Document(NamaTemplate)
    for table in document.tables: #for pertama
        for row in table.rows: #for kedua
            for cell in row.cells: #for ketiga
                for paragraph in cell.paragraphs: #for keempat
                    paragraph.text = paragraph.text.replace("{{nama}}", str(d[0]))
                    paragraph.text = paragraph.text.replace("{{nip}}", str(d[1]))
                    paragraph.text = paragraph.text.replace("{{pg}}", str(d[2]))
                    paragraph.text = paragraph.text.replace("{{jabatan}}", str(d[3]))
                    paragraph.text = paragraph.text.replace("{{instansi}}", str(d[4]))
    document.save(str(idx+1)+" "+str(d[0])+'.docx')


# In[4]:


for d in data:
    idx=data.index(d)
    mengisi(d,idx,NamaTemplate)
    print(d)


# In[52]:


#sekarang pr nya adalah mengidentifikasi font di table biar pas input tulisannya sama
#dan ganti nomor di file template
#btw masih ada 'nan' kalo kolom sumber data kosong


# In[ ]:





# In[ ]:




